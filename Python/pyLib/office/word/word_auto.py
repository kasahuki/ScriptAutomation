from docx import Document
# 库只支持docx 文档

dx = Document()

# 添加文本本质是依靠操作段落做到的
# p1 = dx.add_paragraph('senjay')
# print(type(p1))
# print(dx.paragraphs)
# print(len(dx.paragraphs))
# print(p1.text)
# p1.clear() # 清除这个段落
# print(p1.text)

# p1 = dx.add_paragraph("第一段")
# p2 = dx.add_heading("第二章")
# p1.insert_paragraph_before("第0章")
# dx.add_page_break() # 注意分页符也是段落
# p3 = dx.add_heading("第三章")
# p = p2._element # 先获取标签
# p.getparent().remove(p) # clear方法 也只是删除文本 并没有删除段落也没有删除段落符 
# p._p = p._element = None # 释放内存

p1 = dx.add_paragraph("senjay")
p2 = dx.add_paragraph("senjay")
p3 = dx.add_paragraph("senjay")
p4 = dx.add_paragraph("senjay")
p5 = dx.add_paragraph("senjay")

# p1.alignment = ''   左/右/居中对齐  -- 段落样式
# p1.style = '' 打印这个段落样式
# dx.add_paragraph("text", style = "")
# for paragraph in dx.paragraphs: 遍历所有段落
dx.save('./office/word/test.docx')  # ./ 表示当前目录通常指的是你当前操作的工作目录或者文件所在的目录
